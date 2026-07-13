"""
Universal Document Converter - Backend Prototype
==================================================
Demonstrates the core architecture for a "convert any format to any format" app:
  1. File comes in -> we detect its category (image / document)
  2. We route it to the right conversion engine for that category
  3. Result is saved and a download link is returned

This prototype implements REAL conversions using free, open libraries:
  - Images (png/jpg/jpeg/webp/bmp/gif/tiff) -> any of those formats        [Pillow]
  - PDF -> plain text                                                      [PyMuPDF]
  - PDF -> images (one PNG per page, zipped)                               [PyMuPDF]
  - DOCX -> plain text                                                     [python-docx]
  - DOCX -> PDF (simple text reflow, not pixel-perfect layout)             [python-docx + reportlab]
  - TXT -> PDF                                                             [reportlab]
  - PDF -> DOCX (structural rebuild: headings/lists/tables/images)         [scripts/pdf-to-docx.js, docx-js]

For production-grade DOCX/PPTX/XLSX <-> PDF conversion with full fidelity
(fonts, layout, images, tables preserved exactly), swap in LibreOffice
headless (`soffice --convert-to`) as described in the README. This prototype
avoids that dependency so it runs anywhere without a heavy install.
"""

import os
import io
import uuid
import zipfile
import subprocess
from pathlib import Path

import tempfile
from flask import Flask, request, send_file, jsonify, send_from_directory
from flask_cors import CORS
from PIL import Image
import fitz  # PyMuPDF
import docx
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.units import inch
from ai_fallback import ai_convert

app = Flask(__name__, static_folder="../frontend", static_url_path="")
CORS(app)

UPLOAD_DIR = Path(tempfile.gettempdir()) / "converter_uploads"
OUTPUT_DIR = Path(tempfile.gettempdir()) / "converter_outputs"
UPLOAD_DIR.mkdir(exist_ok=True)
OUTPUT_DIR.mkdir(exist_ok=True)

REPO_ROOT = Path(__file__).resolve().parent.parent
PDF_TO_DOCX_SCRIPT = REPO_ROOT / "scripts" / "pdf-to-docx.js"

IMAGE_FORMATS = {"png", "jpg", "jpeg", "webp", "bmp", "gif", "tiff"}

# Map of (source_ext, target_ext) -> which engine handles it.
# This is the "routing table" — the heart of a multi-format converter.
SUPPORTED_ROUTES = {
    # images -> images (any combination)
    **{(s, t): "image" for s in IMAGE_FORMATS for t in IMAGE_FORMATS if s != t},
    ("pdf", "txt"): "pdf_to_text",
    ("pdf", "png"): "pdf_to_images",
    ("pdf", "docx"): "pdf_to_docx",
    ("docx", "txt"): "docx_to_text",
    ("docx", "pdf"): "docx_to_pdf",
    ("txt", "pdf"): "txt_to_pdf",
}


def ext_of(filename: str) -> str:
    return filename.rsplit(".", 1)[-1].lower() if "." in filename else ""


@app.route("/", methods=["GET"])
def home():
    return app.send_static_file("index.html")


@app.route("/api/formats", methods=["GET"])
def formats():
    """Tell the frontend what conversions are currently possible."""
    routes = sorted([f"{s}->{t}" for (s, t) in SUPPORTED_ROUTES])
    ai_enabled = bool(os.environ.get("GEMINI_API_KEY"))
    return jsonify({"routes": routes, "ai_fallback_enabled": ai_enabled})


@app.route("/api/convert", methods=["POST"])
def convert():
    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    file = request.files["file"]
    target_format = request.form.get("target_format", "").lower().strip()
    if not target_format:
        return jsonify({"error": "No target_format specified"}), 400

    source_ext = ext_of(file.filename)
    route_key = (source_ext, target_format)

    job_id = str(uuid.uuid4())[:8]
    src_path = UPLOAD_DIR / f"{job_id}.{source_ext}"
    file.save(src_path)

    if route_key in SUPPORTED_ROUTES:
        engine = SUPPORTED_ROUTES[route_key]
        try:
            if engine == "image":
                out_path = _convert_image(src_path, job_id, target_format)
            elif engine == "pdf_to_text":
                out_path = _pdf_to_text(src_path, job_id)
            elif engine == "pdf_to_images":
                out_path = _pdf_to_images_zip(src_path, job_id)
            elif engine == "pdf_to_docx":
                out_path = _pdf_to_docx(src_path, job_id)
            elif engine == "docx_to_text":
                out_path = _docx_to_text(src_path, job_id)
            elif engine == "docx_to_pdf":
                out_path = _docx_to_pdf(src_path, job_id)
            elif engine == "txt_to_pdf":
                out_path = _txt_to_pdf(src_path, job_id)
            else:
                return jsonify({"error": "Engine not implemented"}), 500
        except Exception as e:
            return jsonify({"error": f"Conversion failed: {e}"}), 500
        finally:
            src_path.unlink(missing_ok=True)

        return jsonify({
            "job_id": job_id,
            "download_url": f"/api/download/{out_path.name}",
            "method": "builtin"
        })

    # No hardcoded engine for this pair -> try the AI fallback
    if not os.environ.get("GEMINI_API_KEY"):
        src_path.unlink(missing_ok=True)
        return jsonify({
            "error": f"Conversion {source_ext} -> {target_format} not supported, and "
                     f"GEMINI_API_KEY is not set so the AI fallback is unavailable.",
        }), 400

    out_path = OUTPUT_DIR / f"{job_id}.{target_format}"
    success, message = ai_convert(src_path, source_ext, target_format, out_path)
    src_path.unlink(missing_ok=True)

    if not success:
        return jsonify({"error": f"AI fallback failed: {message}"}), 500

    return jsonify({
        "job_id": job_id,
        "download_url": f"/api/download/{out_path.name}",
        "method": "ai_fallback"
    })


@app.route("/api/download/<filename>", methods=["GET"])
def download(filename):
    path = OUTPUT_DIR / filename
    if not path.exists():
        return jsonify({"error": "File not found or expired"}), 404
    return send_file(path, as_attachment=True)


# ---------------- Conversion engines ----------------

def _convert_image(src_path: Path, job_id: str, target_format: str) -> Path:
    img = Image.open(src_path)
    save_format = "JPEG" if target_format in ("jpg", "jpeg") else target_format.upper()
    if save_format == "JPEG" and img.mode in ("RGBA", "P"):
        img = img.convert("RGB")
    out_path = OUTPUT_DIR / f"{job_id}.{target_format}"
    img.save(out_path, format=save_format)
    return out_path


def _pdf_to_text(src_path: Path, job_id: str) -> Path:
    out_path = OUTPUT_DIR / f"{job_id}.txt"
    doc = fitz.open(src_path)
    text = "\n".join(page.get_text() for page in doc)
    out_path.write_text(text, encoding="utf-8")
    return out_path


def _pdf_to_images_zip(src_path: Path, job_id: str) -> Path:
    out_path = OUTPUT_DIR / f"{job_id}.zip"
    doc = fitz.open(src_path)
    with zipfile.ZipFile(out_path, "w") as zf:
        for i, page in enumerate(doc):
            pix = page.get_pixmap(dpi=150)
            img_bytes = pix.tobytes("png")
            zf.writestr(f"page_{i+1}.png", img_bytes)
    return out_path


def _pdf_to_docx(src_path: Path, job_id: str) -> Path:
    """Shells out to the Node CLI (scripts/pdf-to-docx.js), which rebuilds
    the PDF as a real .docx via docx-js rather than doing a lossy raw
    conversion. Requires node plus the Poppler/LibreOffice tools that
    script depends on to be on PATH for this process."""
    out_path = OUTPUT_DIR / f"{job_id}.docx"
    result = subprocess.run(
        ["node", str(PDF_TO_DOCX_SCRIPT), str(src_path), str(out_path)],
        capture_output=True, text=True, timeout=180,
    )
    if result.returncode != 0:
        detail = (result.stderr or result.stdout or "unknown error").strip()
        raise RuntimeError(f"pdf-to-docx.js failed: {detail}")
    return out_path


def _docx_to_text(src_path: Path, job_id: str) -> Path:
    out_path = OUTPUT_DIR / f"{job_id}.txt"
    d = docx.Document(str(src_path))
    text = "\n".join(p.text for p in d.paragraphs)
    out_path.write_text(text, encoding="utf-8")
    return out_path


def _docx_to_pdf(src_path: Path, job_id: str) -> Path:
    out_path = OUTPUT_DIR / f"{job_id}.pdf"
    d = docx.Document(str(src_path))
    c = canvas.Canvas(str(out_path), pagesize=letter)
    width, height = letter
    y = height - inch
    for p in d.paragraphs:
        line = p.text
        if not line.strip():
            y -= 14
            continue
        # naive word wrap at ~95 chars
        for chunk_start in range(0, len(line), 95):
            chunk = line[chunk_start:chunk_start + 95]
            c.drawString(inch, y, chunk)
            y -= 14
            if y < inch:
                c.showPage()
                y = height - inch
    c.save()
    return out_path


def _txt_to_pdf(src_path: Path, job_id: str) -> Path:
    out_path = OUTPUT_DIR / f"{job_id}.pdf"
    text = src_path.read_text(encoding="utf-8", errors="ignore")
    c = canvas.Canvas(str(out_path), pagesize=letter)
    width, height = letter
    y = height - inch
    for line in text.splitlines() or [""]:
        for chunk_start in range(0, max(len(line), 1), 95):
            chunk = line[chunk_start:chunk_start + 95]
            c.drawString(inch, y, chunk)
            y -= 14
            if y < inch:
                c.showPage()
                y = height - inch
    c.save()
    return out_path


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5050))
    app.run(host="0.0.0.0", port=port, debug=False)